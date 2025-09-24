
# Path to your template document
import os
import sys
import csv
import re
from docx import Document
from docx2pdf import convert
import datetime
from docx.shared import RGBColor
if getattr(sys, 'frozen', False):
    # If frozen, sys._MEIPASS is the path to the temporary folder where files are extracted
    application_path = sys._MEIPASS
else:
    # If not frozen, it's running as a normal Python script
    application_path = os.path.dirname(os.path.abspath(__file__))
PDF_TEMPLATE_PATH = os.path.join(application_path, "inputs", "EN_WERFY_Corporate Proposal_[Company name].docx")
CSV_FILE_PATH = os.path.join(application_path, "Montreal Market Data 2025(Sheet1).csv")



def main():
    templatePath = initCreatePDFs()
    generatePDFs(templatePath)

def initCreatePDFs():
    template_path = PDF_TEMPLATE_PATH

    os.makedirs("./attachements", exist_ok=True)
    return template_path

def generatePDFs(template_path):

    with open(CSV_FILE_PATH, "r", newline='', encoding='utf-8') as data:
        csv_reader = csv.DictReader(data)
        now = datetime.datetime.now()
        formattedDate = now.strftime("%d %B %Y")
        all_rows = list(csv_reader) 
        i = 0
        max_rows = len(all_rows)
        print(max_rows)
        for row in all_rows:
            # Extract variables
            recipientName = row["Name of recipient"]
            companyName = row["Company name"]
            email = row["Email"]
            date = formattedDate

            variables = {
                "[name of recipient]": recipientName,
                "[company name]": companyName,
                "[email]": email,
                "[11 JULY 2025]": date
            }
            doc = Document(template_path)
            doc = replaceVariables(doc, variables)


            output_docx_path = f"./attachements/EN_WERFY_Corporate Proposal_{companyName}.docx"
            output_pdf_path = f"./attachements/EN_WERFY_Corporate Proposal_{companyName}.pdf"

            saveDocAsPdf(doc,output_docx_path,output_pdf_path)
            print(f"\n{i}/{max_rows} completed\n")
            i+=1
            # save as docX

        print("\nAll PDFs generated and temporary DOCX files removed.")

def replaceVariables(doc, variables):

    # Replace in paragraphs
    for para in doc.paragraphs:
        replaceTextInPDF(para, variables)

    # Replace in tables
    for table in doc.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                for para in cell.paragraphs:
                    replaceTextInPDF(para, variables)

    # Replace in headers and footers
    for section in doc.sections:
        # Header
        header = section.header
        for para in header.paragraphs:
            replaceTextInPDF(para, variables)
        
        # Footer
        footer = section.footer
        for para in footer.paragraphs:
            replaceTextInPDF(para, variables)
    return doc

# output pdf path is optinal, if it is given, it will be saved as pdf, if not just as docX
def saveDocAsPdf(doc,output_docx_path, output_pdf_path):
    doc.save(output_docx_path)
    # if no pdf path is defined return
    if not output_pdf_path:
        return
    try:
        # convert to pdf
        convert(output_docx_path, output_pdf_path)
    except Exception as e:
        print(f"\nError converting {output_docx_path} to PDF: {e}")
    finally:
        # delete docX temporary file
        if os.path.exists(output_docx_path):
            os.remove(output_docx_path)
def replaceTextInPDF(paragraph, variables):
    full_text = paragraph.text
    
    # Check if any replacement is needed
    replacement_made = False
    for key, value in variables.items():
        if key.lower() in full_text.lower():
            full_text = re.sub(re.escape(key), value, full_text, flags=re.IGNORECASE)
            replacement_made = True
    
    if not replacement_made:
        return
    # Store the formatting of the first run
    first_run_format = None
    if not paragraph.runs:
        return
    first_run = paragraph.runs[0]
    first_run_format = {
        'bold': first_run.bold,
        'italic': first_run.italic,
        'underline': first_run.underline,
        'font_name': first_run.font.name,
        'font_size': first_run.font.size
    }
    
    # Clear all runs
    for run in paragraph.runs:
        run.text = ""
    
    # Create new run with replaced text
    if paragraph.runs:
        new_run = paragraph.runs[0]
    else:
        new_run = paragraph.add_run()
    
    new_run.text = full_text
    
    # Apply formatting if we saved it
    if not first_run_format:
        return 
    try:
        if first_run_format['bold'] is not None:
            new_run.bold = first_run_format['bold']
        if first_run_format['italic'] is not None:
            new_run.italic = first_run_format['italic']
        if first_run_format['underline'] is not None:
            new_run.underline = first_run_format['underline']
        if first_run_format['font_name']:
            new_run.font.name = first_run_format['font_name']
        if first_run_format['font_size']:
            new_run.font.size = first_run_format['font_size']
        new_run.font.color.rgb = RGBColor(0, 0, 0)
    except:
        pass 


if __name__ == "__main__":
    main()

