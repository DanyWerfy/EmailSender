from docx import Document
import csv
import os
import re
from docx2pdf import convert
import datetime

# Path to your template document
template_path = "./attachements/EN_WERFY_Corporate Proposal_[Company name].docx"

# Ensure the output folder exists
os.makedirs("./output", exist_ok=True)


def replaceText(paragraph, variables):
    """Simple method: reconstruct entire paragraph"""
    full_text = paragraph.text
    
    # Check if any replacement is needed
    replacement_made = False
    for key, value in variables.items():
        if key.lower() in full_text.lower():
            full_text = re.sub(re.escape(key), value, full_text, flags=re.IGNORECASE)
            replacement_made = True
    
    if replacement_made:
        # Store the formatting of the first run (if any)
        first_run_format = None
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
                'font_name': first_run.font.name,
                'font_size': first_run.font.size
            }
        
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""
        
        # Create new run with replaced text
        if paragraph.runs:
            new_run = paragraph.runs[0]
        else:
            new_run = paragraph.add_run()
        
        new_run.text = full_text
        
        # Apply formatting if we saved it
        if first_run_format:
            try:
                if first_run_format['bold'] is not None:
                    new_run.bold = first_run_format['bold']
                if first_run_format['italic'] is not None:
                    new_run.italic = first_run_format['italic']
                if first_run_format['underline'] is not None:
                    new_run.underline = first_run_format['underline']
                if first_run_format['font_name']:
                    new_run.font.name = first_run_format['font_name']
                if first_run_format['font_size']:
                    new_run.font.size = first_run_format['font_size']
            except:
                pass  # If formatting fails, continue without it

# Read CSV data
with open("Montreal Market Data 2025(Sheet1).csv", "r", newline='', encoding='utf-8') as data:
    csv_reader = csv.DictReader(data)
    now = datetime.datetime.now()
    formattedDate = now.strftime("%d %B %Y")

    for row in csv_reader:
        # Extract variables
        recipientName = row["Name of recipient"]
        companyName = row["Company name"]
        email = row["Email"]
        date = formattedDate

        variables = {
            "[name of recipient]": recipientName,
            "[company name]": companyName,
            "[email]": email,
            "[11 JULY 2025]": date
        }

        # Load the Word template
        doc = Document(template_path)

        # Replace in paragraphs
        for para in doc.paragraphs:
            replaceText(para, variables)

        # Replace in tables
        for table in doc.tables:
            for table_row in table.rows:
                for cell in table_row.cells:
                    for para in cell.paragraphs:
                        replaceText(para, variables)

        # Replace in headers and footers
        for section in doc.sections:
            # Header
            header = section.header
            for para in header.paragraphs:
                replaceText(para, variables)
            
            # Footer
            footer = section.footer
            for para in footer.paragraphs:
                replaceText(para, variables)

        output_path = f"./output/EN_WERFY_Corporate Proposal_{companyName}.docx"
        doc.save(output_path)
        # this line will conver the docx into a pdf
        # convert(output_path)
